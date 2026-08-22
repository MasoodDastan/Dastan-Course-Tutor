import os
import glob
import base64
from datetime import date, datetime

import anthropic
import requests
import streamlit as st
from dotenv import load_dotenv
import extra_streamlit_components as stx
from PIL import Image

from config import (COURSE_NAME, END_DATE, INSTRUCTOR,
                    MAX_TOKENS, MODEL, REQUIRE_PASSWORD, RESOURCES_PATH,
                    SEMESTER, START_DATE)
from system_prompt import build_system_prompt

load_dotenv()

CLASS_PASSWORD = os.getenv("CLASS_PASSWORD") or st.secrets.get("CLASS_PASSWORD", "")

# ─── Paths ─────────────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ─── Page Config ───────────────────────────────────────────────────────────────
st.set_page_config(page_title="Maya — QMB 2301 Tutor", page_icon="🌟")

MAYA_AVATAR = Image.open(os.path.join(BASE_DIR, "assets", "maya_clean.png"))
COOKIE_NAME = "maya_auth"


# ─── Helpers ───────────────────────────────────────────────────────────────────
@st.cache_data(show_spinner=False)
def load_resources() -> str:
    token = os.getenv("GITHUB_RESOURCES_TOKEN") or st.secrets.get("GITHUB_RESOURCES_TOKEN", "")
    repo  = os.getenv("RESOURCES_REPO") or st.secrets.get("RESOURCES_REPO", "MasoodDastan/Dastan-Course-Tutor-Resources")

    if token:
        headers  = {"Authorization": f"token {token}", "Accept": "application/vnd.github.v3+json"}
        api_url  = f"https://api.github.com/repos/{repo}/contents/{RESOURCES_PATH}"
        resp     = requests.get(api_url, headers=headers)
        if resp.ok:
            files = sorted([f for f in resp.json() if f["name"].endswith(".txt") and not f["name"].startswith("lecture_")],
                           key=lambda x: x["name"])
            parts = []
            for f in files:
                file_resp = requests.get(f["url"], headers=headers)
                if file_resp.ok:
                    content = base64.b64decode(file_resp.json()["content"]).decode("utf-8").strip()
                    if content:
                        parts.append(f"### {f['name']}\n{content}")
            return "\n\n".join(parts)

    # fallback: local files (for development)
    files = sorted(glob.glob(os.path.join(BASE_DIR, "resources", "*.txt")))
    parts = []
    for path in files:
        with open(path, "r", encoding="utf-8") as fh:
            content = fh.read().strip()
        if content:
            parts.append(f"### {os.path.basename(path)}\n{content}")
    return "\n\n".join(parts)


def is_within_semester() -> bool:
    return START_DATE <= date.today() <= END_DATE


def get_client() -> anthropic.Anthropic:
    api_key = os.getenv("ANTHROPIC_API_KEY") or st.secrets.get("ANTHROPIC_API_KEY")
    if not api_key:
        st.error("ANTHROPIC_API_KEY is not set.")
        st.stop()
    return anthropic.Anthropic(api_key=api_key)


def get_sheet():
    try:
        import gspread
        from google.oauth2.service_account import Credentials
        scopes = ["https://www.googleapis.com/auth/spreadsheets"]
        creds_path = os.path.join(BASE_DIR, "credentials.json")
        if os.path.exists(creds_path):
            creds = Credentials.from_service_account_file(creds_path, scopes=scopes)
        elif "gcp_service_account" in st.secrets:
            creds = Credentials.from_service_account_info(st.secrets["gcp_service_account"], scopes=scopes)
        else:
            return None
        sheet_id = os.getenv("GOOGLE_SHEET_ID") or st.secrets.get("GOOGLE_SHEET_ID")
        if not sheet_id:
            return None
        gc = gspread.authorize(creds)
        ws = gc.open_by_key(sheet_id).sheet1
        if not ws.get_all_values():
            ws.append_row(["Timestamp", "Semester", "Question", "Answer"])
        return ws
    except Exception:
        return None


def log_exchange(question: str, answer: str):
    sheet = get_sheet()
    if sheet is None:
        return
    try:
        sheet.append_row([
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            SEMESTER, question, answer,
        ])
    except Exception:
        pass


def process_upload(uploaded_file):
    """Convert uploaded file into an API content block."""
    if uploaded_file is None:
        return None
    name = uploaded_file.name.lower()
    if name.endswith((".png", ".jpg", ".jpeg")):
        raw  = uploaded_file.read()
        data = base64.standard_b64encode(raw).decode("utf-8")
        mime = "image/png" if name.endswith(".png") else "image/jpeg"
        return {"type": "image", "source": {"type": "base64", "media_type": mime, "data": data}}
    if name.endswith((".xlsx", ".xls", ".csv")):
        import pandas as pd
        df = pd.read_excel(uploaded_file) if not name.endswith(".csv") else pd.read_csv(uploaded_file)
        table = df.to_markdown(index=False)
        return {"type": "text", "text": f"[Uploaded file: {uploaded_file.name}]\n\n{table}"}
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(uploaded_file)
        text = "\n\n".join((page.extract_text() or "") for page in reader.pages).strip()
        if not text:
            text = "(No extractable text found in this PDF — it may be a scan. Try uploading a screenshot instead.)"
        return {"type": "text", "text": f"[Uploaded file: {uploaded_file.name}]\n\n{text}"}
    if name.endswith(".docx"):
        import docx
        document = docx.Document(uploaded_file)
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()
        if not text:
            text = "(No extractable text found in this document.)"
        return {"type": "text", "text": f"[Uploaded file: {uploaded_file.name}]\n\n{text}"}
    return None


def display_content(content):
    """Render a message's content — handles both plain strings and block lists."""
    if isinstance(content, str):
        st.markdown(content)
    else:
        for block in content:
            if block["type"] == "text":
                st.markdown(block["text"])
            elif block["type"] == "image":
                img_bytes = base64.b64decode(block["source"]["data"])
                st.image(img_bytes)


def text_of(content) -> str:
    """Extract plain text from content for logging."""
    if isinstance(content, str):
        return content
    return " ".join(b["text"] for b in content if b["type"] == "text")


# ─── Session state init ────────────────────────────────────────────────────────
cookie_manager = stx.CookieManager()
auth_cookie = cookie_manager.get(COOKIE_NAME)

if "authenticated" not in st.session_state:
    st.session_state.authenticated = (not REQUIRE_PASSWORD) or (auth_cookie == CLASS_PASSWORD)
if "messages" not in st.session_state:
    st.session_state.messages = []
if "uploader_key" not in st.session_state:
    st.session_state.uploader_key = 0


# ─── Date Gate ─────────────────────────────────────────────────────────────────
if not is_within_semester():
    st.title("🌟 Maya")
    st.caption(f"{COURSE_NAME} · {INSTRUCTOR}")
    st.info(
        f"Maya is available from **{START_DATE.strftime('%B %d, %Y')}** "
        f"to **{END_DATE.strftime('%B %d, %Y')}** ({SEMESTER}).\n\n"
        "Please check back when the semester begins."
    )
    st.stop()


# ─── Password Gate ─────────────────────────────────────────────────────────────
if REQUIRE_PASSWORD and not st.session_state.authenticated:
    st.title("Hi, I'm Maya! 👋")
    st.caption(f"{COURSE_NAME}")
    st.write("Enter the class password to get started.")
    with st.form("login"):
        pw = st.text_input("Password", type="password")
        submitted = st.form_submit_button("Enter")
    if submitted:
        if pw == CLASS_PASSWORD:
            cookie_manager.set(
                COOKIE_NAME, CLASS_PASSWORD,
                expires_at=datetime(END_DATE.year, END_DATE.month, END_DATE.day)
            )
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("Incorrect password. Please try again.")
    st.stop()


# ─── Main Chat UI ──────────────────────────────────────────────────────────────
st.title("🌟 Maya")
st.caption(f"{COURSE_NAME} · {INSTRUCTOR}")

with st.expander("ℹ️ How Maya works", expanded=False):
    st.markdown(
        "**What Maya is good at:**\n"
        "- Answering questions using Dr. Dastan's actual lecture content and examples\n"
        "- Quick policy questions — due dates, retake rules, grade breakdown\n"
        "- Generating practice problems and quizzing you on chapters\n"
        "- Walking through problems with explanations, not just answers\n"
        "- Analyzing screenshots, Excel files, or PDFs/Word docs you upload (use the sidebar) — e.g., your syllabus\n\n"
        "**Where Maya has limits:**\n"
        "- She only knows what Dr. Dastan has shared — not the entire internet\n"
        "- She won't complete graded assignments or exams for you\n"
        "- For complex questions outside the course, tools like ChatGPT or Claude may serve you better\n\n"
        "Think of Maya as your **first stop** — fast, course-specific, always available. "
        "For everything else, Dr. Dastan's office hours are MTWR 2:00–3:00 pm, Room 227, College of Business (COBA).\n\n"
        "*Conversations are logged anonymously to help improve Maya over time.*"
    )

# Greeting on fresh session
if not st.session_state.messages:
    with st.chat_message("assistant", avatar=MAYA_AVATAR):
        st.markdown(
            "¡Hola! I'm Maya — a tutor built specifically for QMB 2301. "
            "I know Dr. Dastan's lectures, your syllabus, and the course material inside out.\n\n"
            "I'm your best resource for **course-specific questions** — concepts from the lectures, "
            "practice problems, due dates, or policy questions. "
            "For anything outside this course, other AI tools might serve you better.\n\n"
            "You can also attach a **screenshot, Excel file, or a PDF/Word document** (like your syllabus) from the sidebar. "
            "What do you need help with? 😊"
        )

# Chat history
for msg in st.session_state.messages:
    avatar = MAYA_AVATAR if msg["role"] == "assistant" else None
    with st.chat_message(msg["role"], avatar=avatar):
        display_content(msg["content"])

# Chat input
if user_input := st.chat_input("Ask Maya a question…"):
    upload_block = process_upload(st.session_state.get("pending_upload"))

    if upload_block:
        content = [upload_block, {"type": "text", "text": user_input}]
        st.session_state.uploader_key += 1   # resets the file uploader
        st.session_state.pending_upload = None
    else:
        content = user_input

    st.session_state.messages.append({"role": "user", "content": content})
    with st.chat_message("user"):
        display_content(content)

    with st.chat_message("assistant", avatar=MAYA_AVATAR):
        with st.spinner("Maya is thinking…"):
            try:
                client = get_client()
                system = build_system_prompt(load_resources())
                response = client.messages.create(
                    model=MODEL,
                    max_tokens=MAX_TOKENS,
                    system=[{"type": "text", "text": system, "cache_control": {"type": "ephemeral"}}],
                    messages=st.session_state.messages,
                )
                reply = response.content[0].text
            except anthropic.RateLimitError:
                reply = "Maya is a little overwhelmed right now — too many requests at once. Please wait a moment and try again."
            except anthropic.APIStatusError as e:
                reply = f"Something went wrong on Maya's end (error {e.status_code}). Please try again in a moment."
            except Exception:
                reply = "Maya ran into an unexpected error. Please try again, or contact Dr. Dastan if the problem continues."
        st.markdown(reply)

    st.session_state.messages.append({"role": "assistant", "content": reply})
    log_exchange(text_of(content), reply)

# Sidebar
with st.sidebar:
    st.markdown("### 🌟 Maya")
    st.caption(f"Messages this session: {len(st.session_state.messages)}")
    if st.button("Clear conversation"):
        st.session_state.messages = []
        st.rerun()
    st.divider()

    st.markdown("**📎 Attach a file**")
    st.caption("Screenshot (PNG/JPG), spreadsheet (Excel/CSV), or document (PDF/Word)")
    uploaded = st.file_uploader(
        "Upload",
        type=["png", "jpg", "jpeg", "xlsx", "xls", "csv", "pdf", "docx"],
        label_visibility="collapsed",
        key=f"uploader_{st.session_state.uploader_key}",
    )
    if uploaded:
        st.session_state.pending_upload = uploaded
        st.success(f"✓ {uploaded.name} ready — type your question and send")

    st.divider()
    st.caption(f"Instructor: {INSTRUCTOR}")
    st.caption("sdastan@utep.edu")
